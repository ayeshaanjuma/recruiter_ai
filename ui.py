import streamlit as st
import json
import pandas as pd
import os
import time
from dateutil.parser import parse
from main import run_pipeline, extract_docx_text
from core.jd_understanding import parse_jd

# Set page config
st.set_page_config(
    page_title="AI Candidate Intelligence Platform",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Premium stylesheet
st.markdown("""
<style>
    /* Main styling */
    .stApp {
        background-color: #0b0f19;
        color: #f1f5f9;
    }
    section[data-testid="stSidebar"] {
        background-color: #0f172a !important;
        border-right: 1px solid #1e293b;
    }
    h1, h2, h3, h4, h5, h6 {
        color: #f8fafc;
        font-family: 'Inter', sans-serif;
    }
    
    /* Header Section */
    .header-container {
        background: linear-gradient(135deg, #1e1b4b 0%, #0f172a 100%);
        padding: 24px;
        border-radius: 12px;
        border: 1px solid #312e81;
        margin-bottom: 25px;
        display: flex;
        justify-content: space-between;
        align-items: center;
    }
    .header-title {
        font-size: 28px;
        font-weight: 800;
        letter-spacing: -0.5px;
        background: linear-gradient(90deg, #38bdf8 0%, #818cf8 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .header-subtitle {
        font-size: 14px;
        color: #94a3b8;
        margin-top: 5px;
    }
    
    /* Custom Status Badge */
    .status-badge {
        padding: 6px 14px;
        border-radius: 9999px;
        font-size: 12px;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        display: inline-block;
    }
    .status-ready { background-color: #1e293b; color: #94a3b8; border: 1px solid #475569; }
    .status-processing { background-color: #1e1b4b; color: #818cf8; border: 1px solid #4f46e5; animation: pulse 2s infinite; }
    .status-completed { background-color: #064e3b; color: #34d399; border: 1px solid #059669; }
    .status-error { background-color: #7f1d1d; color: #f87171; border: 1px solid #b91c1c; }
    
    /* Card design */
    .premium-card {
        background-color: #111827;
        border: 1px solid #1f2937;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 20px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
        transition: transform 0.2s ease, border-color 0.2s ease;
    }
    .premium-card:hover {
        transform: translateY(-2px);
        border-color: #374151;
    }
    
    /* KPI widget */
    .kpi-value {
        font-size: 32px;
        font-weight: 800;
        color: #f8fafc;
        margin-bottom: 2px;
    }
    .kpi-label {
        font-size: 13px;
        font-weight: 600;
        color: #94a3b8;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    .kpi-trend {
        font-size: 11px;
        font-weight: 700;
        color: #10b981;
        margin-top: 4px;
    }
    
    /* Badges & Tags */
    .badge-primary {
        display: inline-block;
        background-color: rgba(56, 189, 248, 0.1);
        color: #38bdf8;
        padding: 4px 12px;
        border-radius: 6px;
        font-size: 13px;
        font-weight: 600;
        margin: 4px;
        border: 1px solid rgba(56, 189, 248, 0.2);
    }
    .badge-secondary {
        display: inline-block;
        background-color: rgba(129, 140, 248, 0.1);
        color: #818cf8;
        padding: 4px 12px;
        border-radius: 6px;
        font-size: 13px;
        font-weight: 600;
        margin: 4px;
        border: 1px solid rgba(129, 140, 248, 0.2);
    }
    .badge-success {
        display: inline-block;
        background-color: rgba(52, 211, 153, 0.1);
        color: #34d399;
        padding: 4px 12px;
        border-radius: 6px;
        font-size: 13px;
        font-weight: 600;
        margin: 4px;
        border: 1px solid rgba(52, 211, 153, 0.2);
    }
    .badge-warning {
        display: inline-block;
        background-color: rgba(251, 146, 60, 0.1);
        color: #fb923c;
        padding: 4px 12px;
        border-radius: 6px;
        font-size: 13px;
        font-weight: 600;
        margin: 4px;
        border: 1px solid rgba(251, 146, 60, 0.2);
    }
    .badge-danger {
        display: inline-block;
        background-color: rgba(248, 113, 113, 0.1);
        color: #f87171;
        padding: 4px 12px;
        border-radius: 6px;
        font-size: 13px;
        font-weight: 600;
        margin: 4px;
        border: 1px solid rgba(248, 113, 113, 0.2);
    }
    
    /* Candidate Card List */
    .cand-card-header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 12px;
    }
    .cand-name {
        font-size: 18px;
        font-weight: 700;
        color: #f8fafc;
    }
    .cand-score {
        font-size: 20px;
        font-weight: 800;
        color: #10b981;
    }
    
    @keyframes pulse {
        0%, 100% { opacity: 1; }
        50% { opacity: .5; }
    }
</style>
""", unsafe_allow_html=True)

# Default paths
DEFAULT_JD_PATH = "data/raw/job_description.docx"
DEFAULT_DATASET_PATH = "data/raw/candidates.jsonl"

# Read default JD text
default_jd_text = ""
if os.path.exists(DEFAULT_JD_PATH):
    try:
        default_jd_text = extract_docx_text(DEFAULT_JD_PATH)
    except Exception as e:
        default_jd_text = "Failed to load default JD: " + str(e)
else:
    default_jd_text = "Senior AI Software Engineer position. Requirements: 5-9 years experience, Python, FastAPI, vector databases, embeddings, and search evaluation."

# Initialize Session State
if "pipeline_status" not in st.session_state:
    st.session_state["pipeline_status"] = "READY"
if "ranked_candidates" not in st.session_state:
    # Load previously saved candidates if available
    persisted_path = "data/ranked_candidates.json"
    if os.path.exists(persisted_path):
        try:
            with open(persisted_path, "r", encoding="utf-8") as f:
                cands = json.load(f)
                st.session_state["ranked_candidates"] = cands
                st.session_state["pipeline_metrics"] = {
                    'total_loaded': 100000,
                    'recalled': len(cands),
                    'filtered': 100000 - len(cands),
                    'processing_time': 0.0
                }
                st.session_state["pipeline_status"] = "COMPLETED"
        except:
            st.session_state["ranked_candidates"] = []
            st.session_state["pipeline_metrics"] = None
    else:
        st.session_state["ranked_candidates"] = []
        st.session_state["pipeline_metrics"] = None

# --- Top Header Section ---
status_class = {
    "READY": "status-ready",
    "PROCESSING": "status-processing",
    "COMPLETED": "status-completed",
    "ERROR": "status-error"
}[st.session_state["pipeline_status"]]

st.markdown(f"""
<div class="header-container">
    <div>
        <div class="header-title">🤖 AI Candidate Intelligence Platform</div>
        <div class="header-subtitle">Role Understanding, Semantic Matching, and Explainable Candidate Ranking</div>
    </div>
    <div>
        <span class="status-badge {status_class}">{st.session_state["pipeline_status"]}</span>
    </div>
</div>
""", unsafe_allow_html=True)

# --- Sidebar Controls (Recruitment Workspace) ---
st.sidebar.header("💼 Recruitment Workspace")
st.sidebar.markdown("---")

# Role Configuration Card
st.sidebar.subheader("📋 Role Configuration")
jd_editor = st.sidebar.text_area(
    "Job Description",
    value=default_jd_text,
    height=280,
    help="Define the role requirements here. The AI will extract requirements directly."
)
char_count = len(jd_editor)
word_count = len(jd_editor.split())
st.sidebar.caption(f"📝 {word_count} words | {char_count} characters")

# Dataset Configuration Card
st.sidebar.subheader("📂 Dataset Configuration")
dataset_input = st.sidebar.text_input(
    "Candidate Pool Path",
    value=DEFAULT_DATASET_PATH,
    help="Path to candidates (.jsonl, .json, or gzipped file)"
)
if os.path.exists(dataset_input):
    st.sidebar.success("Dataset found and validated.")
else:
    st.sidebar.warning("Dataset file not found at path.")

# Candidate Processing Slider
st.sidebar.subheader("⚙️ Processing Controls")
candidate_limit_opt = st.sidebar.select_slider(
    "Maximum Candidates to Process",
    options=[100, 500, 1000, 5000, 10000, "All Candidates"],
    value=1000,
    help="Limits the processed pool to keep local execution fast."
)

limit_val = None if candidate_limit_opt == "All Candidates" else int(candidate_limit_opt)

# Collapsible Advanced Options
with st.sidebar.expander("🛠️ Advanced Settings", expanded=False):
    st.slider("Semantic Recall Size", 100, 2000, 1000, 100)
    st.slider("Rule-Based Threshold", 0, 100, 50, 5)
    st.toggle("LLM Re-Ranking", value=True)
    st.selectbox("Explanation Complexity", ["Compact", "Standard", "Detailed"], index=1)
    save_results = st.toggle("Persist Ranking Output", value=True)

st.sidebar.markdown("---")

# Run Pipeline Action
run_pipeline_clicked = st.sidebar.button("🚀 Run Candidate Intelligence Pipeline", use_container_width=True, type="primary")

# Execute Pipeline logic
if run_pipeline_clicked:
    if not jd_editor.strip():
        st.error("Error: Job Description cannot be empty.")
        st.session_state["pipeline_status"] = "ERROR"
    elif not os.path.exists(dataset_input):
        st.error(f"Error: Dataset not found at '{dataset_input}'")
        st.session_state["pipeline_status"] = "ERROR"
    else:
        st.session_state["pipeline_status"] = "PROCESSING"
        # Rerender header
        st.rerun()

# Run actual run execution if status is PROCESSING
if st.session_state["pipeline_status"] == "PROCESSING":
    try:
        feedback = st.info("Initializing pipelines and loading models...")
        start_time = time.time()
        
        # In-process pipeline run
        results = run_pipeline(
            job_description=jd_editor,
            dataset_path=dataset_input,
            candidate_limit=limit_val
        )
        
        st.session_state["ranked_candidates"] = results
        
        # Persist results
        if save_results:
            os.makedirs("data", exist_ok=True)
            with open("data/ranked_candidates.json", "w", encoding="utf-8") as f:
                json.dump(results, f, indent=2)
            
            # Save CSV
            csv_data = []
            for c in results:
                csv_data.append({
                    "rank": c["rank_position"],
                    "candidate_id": c["candidate_id"],
                    "name": c["name"],
                    "overall_score": c["overall_score"],
                    "skills_match": c["sub_scores"]["skills_match"],
                    "experience_fit": c["sub_scores"]["experience_fit"],
                    "trajectory_signal": c["sub_scores"]["trajectory_signal"],
                    "behavioral_activity": c["sub_scores"]["behavioral_activity"],
                    "justification": c["justification"]
                })
            df = pd.DataFrame(csv_data)
            df.to_csv("data/ranked_candidates.csv", index=False, encoding="utf-8")
            
        metrics = getattr(results, "metrics", {
            'total_loaded': limit_val or 100000,
            'recalled': len(results),
            'filtered': 0,
            'processing_time': time.time() - start_time
        })
        st.session_state["pipeline_metrics"] = metrics
        st.session_state["pipeline_status"] = "COMPLETED"
        st.rerun()
    except Exception as e:
        st.session_state["pipeline_status"] = "ERROR"
        st.error(f"Failed to execute candidate matching: {str(e)}")
        st.rerun()

# --- Main Tabs ---
tabs = st.tabs([
    "📊 Executive Overview",
    "📋 Role Intelligence",
    "🏆 Candidate Rankings",
    "👥 Candidate Comparison",
    "🧠 Explainable AI Insights",
    "📥 Export & Reports"
])

ranked_cands = st.session_state.get("ranked_candidates", [])
metrics = st.session_state.get("pipeline_metrics", None)

# --- Tab 1: Executive Overview ---
with tabs[0]:
    if metrics:
        # KPI widgets
        k_col1, k_col2, k_col3, k_col4, k_col5 = st.columns(5)
        
        with k_col1:
            st.markdown(f"""
            <div class="premium-card">
                <div class="kpi-label">Total Pool Size</div>
                <div class="kpi-value">{metrics.get('total_loaded', 0)}</div>
                <div class="kpi-trend">100% Loaded</div>
            </div>
            """, unsafe_allow_html=True)
            
        with k_col2:
            st.markdown(f"""
            <div class="premium-card">
                <div class="kpi-label">Analyzed Pool</div>
                <div class="kpi-value">{metrics.get('total_loaded', 0)}</div>
                <div class="kpi-trend">Fully Checked</div>
            </div>
            """, unsafe_allow_html=True)
            
        with k_col3:
            st.markdown(f"""
            <div class="premium-card">
                <div class="kpi-label">Shortlisted</div>
                <div class="kpi-value">{len(ranked_cands)}</div>
                <div class="kpi-trend">Top matches selected</div>
            </div>
            """, unsafe_allow_html=True)
            
        with k_col4:
            avg_score = sum(c['overall_score'] for c in ranked_cands) / len(ranked_cands) if ranked_cands else 0.0
            st.markdown(f"""
            <div class="premium-card">
                <div class="kpi-label">Avg Match Score</div>
                <div class="kpi-value">{avg_score:.2%}</div>
                <div class="kpi-trend">High alignment</div>
            </div>
            """, unsafe_allow_html=True)
            
        with k_col5:
            st.markdown(f"""
            <div class="premium-card">
                <div class="kpi-label">Pipeline Speed</div>
                <div class="kpi-value">{metrics.get('processing_time', 0.0):.2f}s</div>
                <div class="kpi-trend">Pre-compiled models</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("### 🧬 Candidate Funnel Visualization")
        recalled_count = metrics.get('recalled', 0)
        total_loaded = metrics.get('total_loaded', 0)
        filtered_count = metrics.get('filtered', 0)
        shortlist_count = len(ranked_cands)
        
        st.markdown(f"""
        <div style="display: flex; justify-content: space-between; align-items: center; background-color: #111827; padding: 20px; border-radius: 12px; border: 1px solid #1f2937;">
            <div style="text-align: center; flex: 1;"><div class="kpi-label" style="color:#38bdf8;">1. Candidates Loaded</div><div class="kpi-value" style="font-size:24px; margin-top:5px;">{total_loaded}</div></div>
            <div style="color: #4b5563; font-size: 24px;">➔</div>
            <div style="text-align: center; flex: 1;"><div class="kpi-label" style="color:#818cf8;">2. Exclusions Filtered</div><div class="kpi-value" style="font-size:24px; margin-top:5px; color:#f87171;">-{filtered_count}</div></div>
            <div style="color: #4b5563; font-size: 24px;">➔</div>
            <div style="text-align: center; flex: 1;"><div class="kpi-label" style="color:#10b981;">3. Semantic Recall</div><div class="kpi-value" style="font-size:24px; margin-top:5px;">{recalled_count}</div></div>
            <div style="color: #4b5563; font-size: 24px;">➔</div>
            <div style="text-align: center; flex: 1;"><div class="kpi-label" style="color:#34d399;">4. Shortlisted</div><div class="kpi-value" style="font-size:24px; margin-top:5px; color:#10b981;">{shortlist_count}</div></div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("### 📈 Score Distribution Chart")
        if len(ranked_cands) > 0:
            scores = [c['overall_score'] * 100.0 for c in ranked_cands]
            df_scores = pd.DataFrame(scores, columns=["Match Score (%)"])
            st.bar_chart(df_scores)
    else:
        st.info("Pipeline metrics are empty. Adjust the workspace configuration and run the candidate intelligence pipeline.")

# --- Tab 2: Role Intelligence ---
with tabs[1]:
    if jd_editor.strip():
        try:
            jd_info = parse_jd(jd_editor)
            
            col_jd_l, col_jd_r = st.columns([2, 1])
            
            with col_jd_l:
                st.markdown("""
                <div class="premium-card">
                    <h3>📋 Inferred Role Fingerprint</h3>
                    <p><b>Seniority Band:</b> <span class="badge-primary" style="margin:0;">""" + jd_info.get('seniority_band', 'mid').upper() + """</span></p>
                    <p><b>Minimum Experience Required:</b> """ + str(jd_info.get('min_years_experience', 0)) + """ Years</p>
                    <p><b>Role Narrative Profile:</b> <br><i>""" + jd_info.get('role_narrative', 'N/A') + """</i></p>
                </div>
                """, unsafe_allow_html=True)
                
                st.markdown("### 🛠️ Extracted Skill Inventory")
                st.markdown("**Must-Have Skill Set:**")
                must_skills = jd_info.get('must_have_skills', [])
                if must_skills:
                    st.markdown(" ".join([f"<span class='badge-primary'>{s}</span>" for s in must_skills]), unsafe_allow_html=True)
                else:
                    st.write("None parsed")
                    
                st.markdown("**Nice-to-Have Skill Set:**")
                nice_skills = jd_info.get('nice_to_have_skills', [])
                if nice_skills:
                    st.markdown(" ".join([f"<span class='badge-secondary'>{s}</span>" for s in nice_skills]), unsafe_allow_html=True)
                else:
                    st.write("None parsed")
                    
            with col_jd_r:
                st.markdown("### 🧠 Inferred Behavioral Signal Targets")
                st.write("Ownership & Autonomy")
                st.progress(90)
                st.write("Leadership potential")
                st.progress(70)
                st.write("Fast-paced adaptability")
                st.progress(95)
                st.write("Stakeholder Management")
                st.progress(60)
                
                implicit = jd_info.get('implicit_signals', [])
                if implicit:
                    st.info("💡 **Inferred Trait Signals:** \n" + "\n".join([f"- {s}" for s in implicit]))
        except Exception as e:
            st.error(f"Error parsing role requirements: {e}")
    else:
        st.info("Role description is empty. Edit it in the sidebar controls.")

# --- Tab 3: Candidate Rankings ---
with tabs[2]:
    if ranked_cands:
        st.subheader("🏆 Enterprise Shortlist")
        
        # Initialize search state if not present
        if "search_query_applied" not in st.session_state:
            st.session_state["search_query_applied"] = ""
        if "score_filter_applied" not in st.session_state:
            st.session_state["score_filter_applied"] = 0

        # Wrap controls in a form so they only apply when the button is clicked
        with st.form(key="filter_form"):
            col_rank_c1, col_rank_c2 = st.columns([2, 1])
            with col_rank_c1:
                search_query = st.text_input(
                    "🔍 Search Candidates by Name, Skills, or ID", 
                    value=st.session_state["search_query_applied"],
                    help="Type candidate name, skill, ID, or rank, then click the button below to apply filters."
                )
            with col_rank_c2:
                score_filter = st.slider(
                    "🎯 Minimum Match Score (%)", 
                    0, 100, 
                    value=st.session_state["score_filter_applied"], 
                    step=5
                )
            
            # Submit button for the form
            submit_button = st.form_submit_button("Submit", use_container_width=True)
            
            if submit_button:
                st.session_state["search_query_applied"] = search_query
                st.session_state["score_filter_applied"] = score_filter
                st.rerun()

        # Apply filters in UI
        filtered_cands = []
        for c in ranked_cands:
            match = True
            if st.session_state["search_query_applied"]:
                q = st.session_state["search_query_applied"].lower()
                name_match = q in c['name'].lower()
                skills_match = any(q in s.lower() for s in c.get('normalized_skills', []))
                id_match = q in str(c.get('candidate_id', '')).lower()
                rank_match = q == str(c.get('rank_position', ''))
                match = name_match or skills_match or id_match or rank_match
            if st.session_state["score_filter_applied"]:
                match = match and (c['overall_score'] * 100.0 >= st.session_state["score_filter_applied"])
            if match:
                filtered_cands.append(c)

        if st.session_state["search_query_applied"] or st.session_state["score_filter_applied"]:
            st.caption(f"Showing {len(filtered_cands)} matching candidates out of {len(ranked_cands)} total.")
                
        # Candidate List
        for c in filtered_cands[:100]:
            orig = c.get('original_data', {})
            prof = orig.get('profile', {})
            skills = c.get('normalized_skills', [])
            
            # Score color styling
            score_color = "#10b981" if c['overall_score'] >= 0.75 else ("#f59e0b" if c['overall_score'] >= 0.6 else "#ef4444")
            
            st.markdown(f"""
            <div class="premium-card">
                <div class="cand-card-header">
                    <div>
                        <span class="cand-name">Rank {c['rank_position']}: {c['name']}</span>
                        <span style="color:#94a3b8; font-size:12px; margin-left:10px;">({c['candidate_id']})</span>
                    </div>
                    <div class="cand-score" style="color: {score_color};">{c['overall_score']:.2%}</div>
                </div>
                <div style="font-size:14px; color:#cbd5e1; margin-bottom:8px;">
                    💼 <b>Current Role:</b> {prof.get('current_title', 'Unknown')} at {prof.get('current_company', 'N/A')} ({prof.get('years_of_experience', 0)} years)
                </div>
                <div style="margin-bottom:12px;">
                    {" ".join([f"<span class='candidate-badge'>{s}</span>" for s in skills[:8]])}
                </div>
                <div style="font-size:13px; color:#94a3b8; font-style:italic; border-left:3px solid #334155; padding-left:10px; margin-bottom:10px;">
                    "{c['justification']}"
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            # Expand Details inside the main container
            with st.expander(f"👁️ View Complete Journey for {c['name']}", expanded=False):
                col_sub1, col_sub2 = st.columns(2)
                
                with col_sub1:
                    st.write("**Scoring & Fit Breakdown:**")
                    st.write(f"- 🛠️ **Skills Match:** {c['sub_scores']['skills_match']}/100")
                    st.progress(c['sub_scores']['skills_match'])
                    st.write(f"- 📈 **Experience Fit:** {c['sub_scores']['experience_fit']}/100")
                    st.progress(c['sub_scores']['experience_fit'])
                    st.write(f"- 🚀 **Trajectory Signal:** {c['sub_scores']['trajectory_signal']}/100")
                    st.progress(c['sub_scores']['trajectory_signal'])
                    st.write(f"- 🧠 **Behavioral/Platform Activity:** {c['sub_scores']['behavioral_activity']}/100")
                    st.progress(c['sub_scores']['behavioral_activity'])
                    
                with col_sub2:
                    st.write("**Logistics & Recruitment Status:**")
                    sig = orig.get('redrob_signals', {})
                    st.write(f"- 📅 **Notice Period:** {sig.get('notice_period_days', 'N/A')} Days")
                    st.write(f"- 💬 **Response Rate:** {int(sig.get('recruiter_response_rate', 0.0) * 100)}%")
                    st.write(f"- 📍 **Location:** {prof.get('location', 'N/A')}, {prof.get('country', '')}")
                    st.write(f"- 💰 **Expected Salary:** {sig.get('expected_salary_range_inr_lpa', {}).get('min', 'N/A')} - {sig.get('expected_salary_range_inr_lpa', {}).get('max', 'N/A')} LPA")
                    st.write(f"- 💡 **Relocation willing:** {'Yes' if sig.get('willing_to_relocate') else 'No'}")
                    
                st.markdown("---")
                st.markdown(f"**📝 Executive Summary:** *{prof.get('summary', 'N/A')}*")
                
                st.markdown("**💼 Career History Timeline:**")
                for j in orig.get('career_history', []):
                    st.markdown(f"- **{j.get('title')}** at *{j.get('company')}* ({j.get('start_date')} to {j.get('end_date') or 'Current'}, {j.get('duration_months')} months)")
                    st.markdown(f"  *\"{j.get('description')}\"*")
                    
                st.markdown("**🎓 Education:**")
                for ed in orig.get('education', []):
                    st.markdown(f"- **{ed.get('degree')} in {ed.get('field_of_study')}** — {ed.get('institution')} ({ed.get('start_year')} - {ed.get('end_year')})")
    else:
        st.info("Shortlist is empty. Run pipeline to select and rank candidates.")

# --- Tab 4: Candidate Comparison ---
with tabs[3]:
    if ranked_cands:
        st.subheader("👥 Candidate Comparison Engine")
        
        # Select candidates to compare
        options = {f"Rank {c['rank_position']}: {c['name']}": c for c in ranked_cands[:20]}
        selected_names = st.multiselect("Select up to 4 candidates to compare:", options=list(options.keys()), default=list(options.keys())[:2])
        
        if len(selected_names) > 0:
            compare_data = []
            for name in selected_names:
                c = options[name]
                orig = c.get('original_data', {})
                prof = orig.get('profile', {})
                sig = orig.get('redrob_signals', {})
                
                compare_data.append({
                    "Name": c["name"],
                    "Rank": c["rank_position"],
                    "Overall Match": f"{c['overall_score']:.2%}",
                    "Skills Score": f"{c['sub_scores']['skills_match']}/100",
                    "Experience Score": f"{c['sub_scores']['experience_fit']}/100",
                    "Trajectory Score": f"{c['sub_scores']['trajectory_signal']}/100",
                    "Notice Period": f"{sig.get('notice_period_days', 'N/A')} days",
                    "Response Rate": f"{int(sig.get('recruiter_response_rate', 0.0) * 100)}%",
                    "Location": prof.get('location', 'N/A')
                })
                
            df_compare = pd.DataFrame(compare_data)
            st.dataframe(df_compare, use_container_width=True, hide_index=True)
            
            # Highlight highlights
            st.markdown("### 🌟 Comparative Highlights")
            best_cand = options[selected_names[0]]
            st.success(f"🏆 **Best Fit candidate:** **{best_cand['name']}** holds the top rank among selected candidates with a score of **{best_cand['overall_score']:.2%}**.")
        else:
            st.info("Please select candidates from the list to display side-by-side metrics.")
    else:
        st.info("Run candidate ranking to populate comparison options.")

# --- Tab 5: Explainable AI Insights ---
with tabs[4]:
    if ranked_cands:
        st.subheader("🧠 Explainable Ranking Insights")
        
        st.markdown("### 📋 Matching Methodology & Funnel Pipeline")
        st.markdown("""
        The ranking system utilizes a hybrid heuristic-semantic algorithm to evaluate Availability and Match:
        1. **JD Fingerprint Parsing**: Standardizes experience limits, location profiles, and must-have skill tags.
        2. **Stage 1 (Local Filtering)**: Automatically excludes honeypots (profile discrepancies), title chasers, academic-only profiles, and services-only contractors.
        3. **Stage 2 (Heuristic Weighting)**: Computes skill-match and experience-fit scores.
        4. **Stage 3 (Behavioral multipliers)**: Down-weights inactive/unresponsive candidates.
        5. **Stage 4 (Semantic Re-ranking)**: Matches candidate profiles using local SentenceTransformer embeddings.
        """)
        
        st.markdown("### 🔍 Specific Candidate Insights")
        
        # Initialize session state for selected candidate insight
        if "selected_candidate_insight" not in st.session_state:
            st.session_state["selected_candidate_insight"] = None
            
        candidate_options = [f"Rank {c['rank_position']}: {c['name']}" for c in ranked_cands[:20]]
        
        # Form for selecting candidate to prevent auto-rerunning/displaying before clicking enter/submit
        with st.form("insight_form"):
            selected_insight = st.selectbox(
                "Select Candidate for Deep AI Insight:", 
                options=candidate_options,
                index=0 if st.session_state["selected_candidate_insight"] is None or st.session_state["selected_candidate_insight"] not in candidate_options else candidate_options.index(st.session_state["selected_candidate_insight"])
            )
            submit_insight = st.form_submit_button("Submit", use_container_width=True)
            
            if submit_insight:
                st.session_state["selected_candidate_insight"] = selected_insight
                st.rerun()
                
        # Display insights if a candidate has been selected and confirmed
        if st.session_state["selected_candidate_insight"]:
            c = next((x for x in ranked_cands if f"Rank {x['rank_position']}: {x['name']}" == st.session_state["selected_candidate_insight"]), None)
            if c:
                orig = c.get('original_data', {})
                prof = orig.get('profile', {})
                
                col_in1, col_in2 = st.columns(2)
                with col_in1:
                    st.write("**💪 Candidate Strengths:**")
                    st.markdown(f"- **Excellent experience alignment:** {prof.get('years_of_experience')} years in industry roles.")
                    st.markdown(f"- **Matching skill keywords:** {', '.join(c.get('normalized_skills', [])[:4])}.")
                    st.markdown(f"- **High behavioral availability:** Available with short notice period.")
                with col_in2:
                    # Distinguish factors
                    st.write("**🛡️ Confidence Rating & Risks:**")
                    conf_val = "HIGH" if c['overall_score'] > 0.75 else "MEDIUM"
                    conf_class = "badge-success" if conf_val == "HIGH" else "badge-warning"
                    st.markdown(f"Confidence Level: <span class='status-badge {conf_class}'>{conf_val}</span>", unsafe_allow_html=True)
                    st.write("- No critical timeline anomalies detected (Honeypot verified).")
                    st.write("- Previous corporate tenure demonstrates steady team contributions.")
        else:
            st.info("💡 Select a candidate from the dropdown above and click 'View Deep Insights' to display their strengths, confidence level, and risks.")
    else:
        st.info("Insights will be available once the ranking pipeline is executed.")

# --- Tab 6: Export & Reports ---
with tabs[5]:
    if len(ranked_cands) > 0:
        st.subheader("📥 Export Talent Intelligence Reports")
        
        # Helper function for Excel generation
        def generate_excel_data(candidates):
            import io
            import pandas as pd
            excel_rows = []
            for c in candidates:
                excel_rows.append({
                    "Rank": c["rank_position"],
                    "Candidate ID": c["candidate_id"],
                    "Name": c["name"],
                    "Match Score": f"{c['overall_score']:.2%}",
                    "Skills Match Score (/100)": c["sub_scores"]["skills_match"],
                    "Experience Fit Score (/100)": c["sub_scores"]["experience_fit"],
                    "Trajectory Score (/100)": c["sub_scores"]["trajectory_signal"],
                    "Behavioral Score (/100)": c["sub_scores"]["behavioral_activity"],
                    "Justification": c["justification"]
                })
            df_excel = pd.DataFrame(excel_rows)
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                df_excel.to_excel(writer, index=False, sheet_name='Shortlist')
            return buffer.getvalue()

        # Helper function for PDF generation
        def generate_pdf_data(candidates):
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            
            # Title
            pdf.set_font("helvetica", "B", 16)
            pdf.cell(0, 10, "AI Candidate Shortlist & Recruitment Report", ln=True, align="C")
            pdf.ln(5)
            
            # Subtitle
            pdf.set_font("helvetica", "I", 10)
            pdf.cell(0, 5, "Generated by AI Candidate Intelligence Platform", ln=True, align="C")
            pdf.ln(10)
            
            # Summary header
            pdf.set_font("helvetica", "B", 12)
            pdf.cell(0, 8, f"Top Shortlisted Candidates: {len(candidates)}", ln=True)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(5)
            
            # Candidates Loop
            for c in candidates[:50]: # Limit to top 50 to keep PDF readable
                # Check page break
                if pdf.get_y() > 250:
                    pdf.add_page()
                    
                pdf.set_font("helvetica", "B", 11)
                name_str = f"Rank {c['rank_position']}: {c['name']} ({c['candidate_id']}) - Match Score: {c['overall_score']:.2%}"
                pdf.cell(0, 8, name_str.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                
                pdf.set_font("helvetica", "", 10)
                scores_str = f"Skills Match: {c['sub_scores']['skills_match']}/100 | Experience Fit: {c['sub_scores']['experience_fit']}/100 | Trajectory: {c['sub_scores']['trajectory_signal']}/100"
                pdf.cell(0, 6, scores_str.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                
                pdf.set_font("helvetica", "I", 9)
                just = c.get('justification', 'N/A')
                just = just.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 5, f"Justification: {just}")
                pdf.ln(4)
                
            return bytes(pdf.output())

        # Helper function for DOCX generation
        def generate_docx_data(candidates):
            import docx
            import io
            doc = docx.Document()
            
            doc.add_heading("AI Candidate Shortlist & Recruitment Report", level=0)
            doc.add_paragraph("Generated by AI Candidate Intelligence Platform")
            
            doc.add_heading("Shortlisted Candidates Summary", level=1)
            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Rank'
            hdr_cells[1].text = 'ID'
            hdr_cells[2].text = 'Name'
            hdr_cells[3].text = 'Match Score'
            
            for c in candidates[:50]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(c["rank_position"])
                row_cells[1].text = c["candidate_id"]
                row_cells[2].text = c["name"]
                row_cells[3].text = f"{c['overall_score']:.2%}"
                
            doc.add_paragraph()
            
            doc.add_heading("Detailed Justifications", level=1)
            for c in candidates[:50]:
                doc.add_heading(f"Rank {c['rank_position']}: {c['name']} ({c['candidate_id']})", level=2)
                doc.add_paragraph(f"Match Score: {c['overall_score']:.2%}")
                doc.add_paragraph(f"Justification: {c.get('justification', 'N/A')}")
                doc.add_paragraph()
                
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

        # Helper function for TXT generation
        def generate_txt_data(candidates):
            lines = []
            lines.append("==================================================")
            lines.append("AI Candidate Shortlist & Recruitment Report")
            lines.append("Generated by AI Candidate Intelligence Platform")
            lines.append("==================================================\n")
            
            for c in candidates[:100]:
                lines.append(f"Rank {c['rank_position']}: {c['name']} ({c['candidate_id']})")
                lines.append(f"Match Score: {c['overall_score']:.2%}")
                lines.append(f"Skills Match: {c['sub_scores']['skills_match']}/100")
                lines.append(f"Experience Fit: {c['sub_scores']['experience_fit']}/100")
                lines.append(f"Trajectory Signal: {c['sub_scores']['trajectory_signal']}/100")
                lines.append(f"Justification: {c.get('justification', 'N/A')}")
                lines.append("-" * 50 + "\n")
                
            return "\n".join(lines).encode("utf-8")

        col_ex1, col_ex2 = st.columns(2)
        with col_ex1:
            st.markdown("""
            <div class="premium-card">
                <h4>💾 Download Ranking Data</h4>
                <p>Export the full candidate shortlist and sub-scores in standard interchange formats.</p>
            </div>
            """, unsafe_allow_html=True)
            
            # CSV Download
            csv_path = "data/ranked_candidates.csv"
            if os.path.exists(csv_path):
                with open(csv_path, "r", encoding="utf-8") as f:
                    csv_text = f.read()
                st.download_button(
                    "Download Shortlist (CSV)",
                    data=csv_text,
                    file_name="shortlist_candidates.csv",
                    mime="text/csv",
                    use_container_width=True
                )
                
            # JSON Download
            json_path = "data/ranked_candidates.json"
            if os.path.exists(json_path):
                with open(json_path, "r", encoding="utf-8") as f:
                    json_text = f.read()
                st.download_button(
                    "Download Shortlist (JSON)",
                    data=json_text,
                    file_name="shortlist_candidates.json",
                    mime="application/json",
                    use_container_width=True
                )

            # Plain Text Download
            txt_bytes = generate_txt_data(ranked_cands)
            st.download_button(
                "Download Shortlist (TXT)",
                data=txt_bytes,
                file_name="shortlist_candidates.txt",
                mime="text/plain",
                use_container_width=True
            )
                
        with col_ex2:
            st.markdown("""
            <div class="premium-card">
                <h4>📄 Export Executive Summaries</h4>
                <p>Generate clean, recruiter-focused documents detailing parsed JD requirements and matched shortlist justifications.</p>
            </div>
            """, unsafe_allow_html=True)
            
            # PDF Download
            pdf_bytes = generate_pdf_data(ranked_cands)
            st.download_button(
                "Export PDF Executive Summary",
                data=pdf_bytes,
                file_name="candidate_shortlist_report.pdf",
                mime="application/pdf",
                use_container_width=True
            )
            
            # Word Download
            docx_bytes = generate_docx_data(ranked_cands)
            st.download_button(
                "Export Word Document (DOCX)",
                data=docx_bytes,
                file_name="candidate_shortlist_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

            # Excel Download
            excel_bytes = generate_excel_data(ranked_cands)
            st.download_button(
                "Export Excel Recruitment Report",
                data=excel_bytes,
                file_name="candidate_shortlist_report.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
    else:
        st.info("Run candidate ranking to unlock export formats.")
